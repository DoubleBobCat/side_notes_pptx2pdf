import os
import shutil
import sys
import cv2
import json
import numpy as np
from pptx_tools import utils
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from win32com import client

dev_mod = False


def convert_pptx2png(folder_path, file_name):
    os.mkdir(os.path.join(folder_path, 'temp_folder\\origin'))
    pptfile = os.path.join(folder_path, file_name)
    png_folder = os.path.join(folder_path, 'temp_folder\\origin')
    utils.save_pptx_as_png(png_folder, pptfile, overwrite_folder=True)
    print('-> done.')


def splice_png(folder_path, pages_in_page):
    # sort image paths
    def sort_image_paths(image_paths):
        return sorted(image_paths, key=lambda x: int(x.split('幻灯片')[1].split('.')[0]))

    # vertical concat
    def vertical_concat(image_paths):
        # get image
        images = [cv2.imread(img_path) for img_path in image_paths]
        # vertical images
        img_temp = images[0]
        if images[0] is not None:
            for img in images:
                if img is not images[0]:
                    img_temp = cv2.vconcat([img_temp, img])
            return img_temp
        else:
            print('Error -> Can not read the image.')
            exit()

    # main flow
    os.mkdir(os.path.join(folder_path, 'temp_folder\\splice'))
    png_folder = os.path.join(folder_path, 'temp_folder\\origin')
    output_folder = os.path.join(folder_path, 'temp_folder\\splice')
    image_paths = []
    sorted_image_paths = []
    output_img_fils = []

    for root, dirs, files in os.walk(png_folder):
        for file in files:
            if file.lower().endswith('.png'):
                image_paths.append(os.path.join(root, file))

    sorted_image_paths_temp = sort_image_paths(image_paths)
    for item in sorted_image_paths_temp:
        os.rename(item, item.replace('幻灯片', ''))
        sorted_image_paths.append(item.replace('幻灯片', ''))

    full_groups, remaining = divmod(len(sorted_image_paths), pages_in_page)

    for i in range(0, full_groups * pages_in_page, pages_in_page):
        concat_img = vertical_concat(sorted_image_paths[i:i+pages_in_page])
        output_filename = os.path.join(
            output_folder, f'{(i//pages_in_page)+1}.PNG')
        cv2.imwrite(output_filename, concat_img)
        output_img_fils.append(output_filename)

    if remaining > 0:
        concat_img = vertical_concat(sorted_image_paths[-remaining:])
        output_filename = os.path.join(
            output_folder, f'{full_groups+1}.PNG')
        cv2.imwrite(output_filename, concat_img)
        output_img_fils.append(output_filename)

    print('-> done.')
    return output_img_fils


def draw_note(output_img_fils, line_num, margin, line_color, line_thickness_):
    # draw a white pic
    page_width = Cm(21) - Cm(margin[2]) - Cm(margin[3])
    page_height = Cm(29.7) - Cm(margin[0]) - Cm(margin[1])
    temple_img = cv2.imread(output_img_fils[0])
    img_height, img_width, _ = temple_img.shape
    multiple = page_width / page_height
    max_width = img_height * multiple
    note_width = max_width - img_width
    white_image = np.full(
        (int(img_height), int(note_width), 3), 255, dtype=np.uint8)
    note_img = white_image
    # draw line
    if line_num != 0:
        interval = img_height // line_num
        line_color_ = (line_color[0], line_color[1], line_color[2])
        line_thickness_ = line_thickness
        for i in range(0, int(img_height), int(interval)):
            cv2.line(note_img, (0, i), (int(note_width) - 1, i),
                     line_color_, line_thickness_)
    # splice img and note
    page = 1
    for images_path in output_img_fils:
        image = cv2.imread(images_path)
        # fill height
        padding_height = note_img.shape[0] - image.shape[0]
        if padding_height != 0:
            padding = np.full(
                (padding_height, image.shape[1], 3), 255, dtype=np.uint8)
            padded_image = np.vstack((image, padding))
            image = padded_image
        # concat
        if double_page_printing:
            if page % 2 == 1:
                image_temp = cv2.hconcat([image, note_img])
            else:
                image_temp = cv2.hconcat([note_img, image])
            page += 1
        else:
            image_temp = cv2.hconcat([image, note_img])
        cv2.imwrite(images_path, image_temp)
    print('-> done.')


def add_png2docx(folder_path, file_name, output_img_fils, margin):
    def create_element(name):
        return OxmlElement(name)

    def create_attribute(element, name, value):
        element.set(qn(name), value)

    def add_page_number(run):
        fldChar1 = create_element('w:fldChar')
        create_attribute(fldChar1, 'w:fldCharType', 'begin')
        instrText = create_element('w:instrText')
        create_attribute(instrText, 'xml:space', 'preserve')
        instrText.text = "PAGE"
        fldChar2 = create_element('w:fldChar')
        create_attribute(fldChar2, 'w:fldCharType', 'end')
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
    picture_weight = Cm(0)
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(margin[0])
    sec.bottom_margin = Cm(margin[1])
    sec.left_margin = Cm(margin[2])
    sec.right_margin = Cm(margin[3])
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    for img_file in output_img_fils:
        picture = doc.add_picture(img_file, height=Cm(29.7-2*margin[0]))
        if picture_weight == Cm(0):
            picture_weight = picture.width.cm
        if picture.width.cm != picture_weight:
            scaling = picture.height.cm / picture.width.cm
            picture.width = Cm(picture_weight)
            picture.height = Cm(picture_weight * scaling)
    if add_page_num:
        add_page_number(doc.sections[0].footer.paragraphs[0].add_run())
        doc.sections[0].footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    output_docx = os.path.join(
        folder_path, 'temp_folder\\' + file_name.replace('pptx', 'docx'))
    doc.save(output_docx)
    print('-> done.')


def convert_docx2pdf(folder_path):
    word = client.Dispatch("Word.Application")
    output_docx = os.path.join(
        folder_path, 'temp_folder\\' + file_name.replace('pptx', 'docx'))
    output_pdf = os.path.join(folder_path, file_name.replace('pptx', 'pdf'))
    doc = word.Documents.Open(output_docx)
    doc.SaveAs(output_pdf, 17)
    doc.Close()
    word.Quit()
    print('-> done.')


if __name__ == "__main__":
    print('----- Programme Begin -----')

    # init var
    folder_path = 'D:\\'
    file_name = 'pptx.pptx'
    pages_in_page = 4
    margin = []
    output_img_fils = []
    double_page_printing = False
    line_num = 32
    add_page_num = False

    # input file
    if dev_mod:
        str = '"D:\School\PPT\古风.pptx"'
    else:
        str = input("Please input file path: ")
    str = str.replace('\"', '')
    str = str.replace('\\', '\\\\')
    folder_path, file_name = os.path.split(str)

    # prepare
    print('-> prepare', end='... ')
    sys.stdout.flush()
    if os.path.isdir(os.path.join(folder_path, 'temp_folder')):
        shutil.rmtree(os.path.join(folder_path, 'temp_folder'))
    if os.path.isfile(file_name.replace('pptx', 'pdf')):
        os.remove(file_name.replace('pptx', 'pdf'))
    os.mkdir(os.path.join(folder_path, 'temp_folder'))
    with open('config.json', 'r') as file:
        config_data = json.load(file)
    double_page_printing = config_data['double_page_printing']
    line_num = config_data['line_num']
    line_color = config_data['line_color']
    line_thickness = config_data['line_thickness']
    pages_in_page = config_data['pages_in_page']
    margin = config_data['margin']
    add_page_num = config_data['add_page_num']
    print('-> done.')

    # convert pptx to png
    print('-> converting PPTX to PNG', end='... ')
    sys.stdout.flush()
    convert_pptx2png(folder_path, file_name)

    # splice png
    print('-> splicing PNG', end='... ')
    sys.stdout.flush()
    output_img_fils = splice_png(folder_path, pages_in_page)

    # draw note
    print('-> draw note', end='... ')
    sys.stdout.flush()
    draw_note(output_img_fils, line_num, margin, line_color, line_thickness)

    # add png to docx
    print('-> add PNG to DOCX', end='... ')
    sys.stdout.flush()
    add_png2docx(folder_path, file_name, output_img_fils, margin)

    # convert docx to pdf
    print('-> converting DOCX to PDF', end='... ')
    sys.stdout.flush()
    convert_docx2pdf(folder_path)

    # end processing
    print('-> end processing', end='... ')
    sys.stdout.flush()
    shutil.rmtree(os.path.join(folder_path, 'temp_folder'))
    print('-> done.')

    # show pdf path
    print('---------------------------')
    print('-> Output PDF to: ' + os.path.join(folder_path,
          file_name.replace('pptx', 'pdf')))

    print('----- Programme Exit ------')
    exit()
